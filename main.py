import os
from queue import Queue
import re

from docx import Document
import numpy as np
import pandas as pd

from helper import *
from student import Student
from settings import Settings, PLACEHOLDER_PREFIX, PLACEHOLDER_SUFFIX


MESSAGE_START_HEADERS_FOOTERS = "Reading headers and footers..."
MESSAGE_START_PARAGRAPHS = "Reading paragraphs..."
MESSAGE_START_TABLES = "Reading tables..."

def replace_in_paragraph(paragraph, dictionary):
    # Eliminate placeholders that are separated into multiple runs.
    paragraph = consolidate_runs(paragraph)
    if is_placeholder_in_text(paragraph.text):
        for run in paragraph.runs:
            for placeholder, replacement in dictionary.items():
                run.text = replace_text(placeholder, replacement, run.text)
    return paragraph

def replace_in_table(table, dictionary, recursive_calls=1):
    for column in table.columns:
        for cell in column.cells:
            if cell.tables and recursive_calls > 0:
                for nested_table in cell.tables:
                    nested_table = replace_in_table(nested_table, dictionary, recursive_calls-1)
            else:
                for paragraph in cell.paragraphs:
                    paragraph = replace_in_paragraph(paragraph, dictionary)
    return table

def main(settings: Settings, queue : Queue = None):
    # List of apostrophe characters to search for. The first character in this list will be used in text replacements.
    APOSTROPHES = ["’", "'"]

    # Open the template.
    document = Document(settings.filename_template)

    # Open the spreadsheet.
    df = pd.read_excel(
        settings.filename_spreadsheet,
        sheet_name=0,  # Get first sheet only
        header=0,  # Get labels from first row
    )
    df = df.add_prefix(PLACEHOLDER_PREFIX)
    df = df.add_suffix(PLACEHOLDER_SUFFIX)
    # Replace any spaces in column headers with underscores to match the text in the template.
    df = df.rename(columns=lambda string: string.replace(" ", "_"))
    # Replace empty cells with None.
    df = df.where(pd.notnull(df), None)
    
    # Store information in a Student object to infer data for blank cells.
    student = Student(
        name_first=df.at[settings.spreadsheet_row, format_as_placeholder("first_name")],
        name_middle=df.at[settings.spreadsheet_row, format_as_placeholder("middle_name")],
        name_last=df.at[settings.spreadsheet_row, format_as_placeholder("last_name")],
        gender=df.at[settings.spreadsheet_row, format_as_placeholder("gender")],
        grade=df.at[settings.spreadsheet_row, format_as_placeholder("grade")],
        age=df.at[settings.spreadsheet_row, format_as_placeholder("age")],
        birthday=df.at[settings.spreadsheet_row, format_as_placeholder("birthday")],
    )

    # Store the data from the spreadsheet in a dictionary, replacing None with empty strings.
    data = {
        label: str(df.at[settings.spreadsheet_row, label]) if df.at[settings.spreadsheet_row, label] is not None else ''
        for label in df.columns
    }
    
    # Insert data inferred from entered data.
    data[format_as_placeholder("full_name")] = student.name_full
    data[format_as_placeholder("he_she")] = student.pronoun_subject
    data[format_as_placeholder("him_her")] = student.pronoun_object
    data[format_as_placeholder("his_her")] = student.pronoun_possessive_dependent
    data[format_as_placeholder("his_hers")] = student.pronoun_possessive_independent
    data[format_as_placeholder("himself_herself")] = student.pronoun_reflexive
    # Check that no header label contains upper case letters.
    assert are_keys_lowercase(data), f"Make sure all headers in {settings.filename_spreadsheet} are lower case and that there are no empty columns."

    # Search for instances of 's following the student's name and omit the "s" if the first name ends with "s".
    if student.name_first[-1].lower() == "s":
        # Insert this item at the beginning so that instances of 's are searched before searching for instances of first name without 's.
        data = {
            **{f"{df.columns[0]}({'|'.join(APOSTROPHES)})s": f"{student.name_first}{APOSTROPHES[0]}"},
            **data,
        }

    # Iterate over the sections' headers and footers in the template.
    print(MESSAGE_START_HEADERS_FOOTERS)
    if queue:
        queue.put(MESSAGE_START_HEADERS_FOOTERS)
    for section in document.sections:
        header_paragraph_count = len(section.header.paragraphs)
        for i, paragraph in enumerate(section.header.paragraphs):
            paragraph = replace_in_paragraph(paragraph, data)
            if queue:
                queue.put(round(100 * i / header_paragraph_count))
        footer_paragraph_count = len(section.footer.paragraphs)
        for i, paragraph in enumerate(section.footer.paragraphs):
            paragraph = replace_in_paragraph(paragraph, data)
            if queue:
                queue.put(round(100 * i / footer_paragraph_count))

    # Iterate over the paragraphs in the template.
    print(MESSAGE_START_PARAGRAPHS)
    if queue:
        queue.put(MESSAGE_START_PARAGRAPHS)
    paragraph_count = len(document.paragraphs)
    for i, paragraph in enumerate(document.paragraphs, 1):
        paragraph = replace_in_paragraph(paragraph, data)
        if queue:
            queue.put(round(100 * i / paragraph_count))

    # Iterate over the tables in the template.
    print(MESSAGE_START_TABLES)
    if queue:
        queue.put(MESSAGE_START_TABLES)
    table_count = len(document.tables)
    for i, table in enumerate(document.tables, 1):
        table = replace_in_table(table, data, settings.nested_tables)
        if queue:
            queue.put(round(100 * i / table_count))
    
    # Save the modified template as a new file.
    document.save(settings.filename_final)
    message_done = f"Wrote {settings.filename_final}."
    print(message_done)
    if queue:
        queue.put(message_done)

if __name__ == "__main__":
    settings = Settings()
    main(settings)